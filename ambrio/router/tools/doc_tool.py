# ambrio/router/tools/doc_tool.py
"""
Document Reading Tool for Ambrio.
Reads PDFs, images, Word docs, Excel files — extracts text for AI processing.

Tools registered:
  doc_read(path)        - Read any document (PDF, image, DOCX, XLSX, TXT)
  doc_extract_table(path) - Extract tables from documents as structured data
"""
import os, logging, subprocess, json
from pathlib import Path
from ambrio.router.tool_registry import tool

log = logging.getLogger(__name__)

# ── Backend detection ─────────────────────────────────────────────────────────
_BACKENDS = {}

try:
    import pdfplumber
    _BACKENDS['pdf'] = 'pdfplumber'
except ImportError:
    pass

try:
    import docx
    _BACKENDS['docx'] = 'python-docx'
except ImportError:
    pass

try:
    import openpyxl
    _BACKENDS['xlsx'] = 'openpyxl'
except ImportError:
    pass

try:
    from PIL import Image
    _BACKENDS['image'] = 'pillow'
except ImportError:
    pass

log.info(f'doc_tool backends: {list(_BACKENDS.keys())}')


def _read_pdf(path: Path) -> str:
    if 'pdf' not in _BACKENDS:
        return f'[PDF reading requires pdfplumber: pip install pdfplumber]'
    import pdfplumber
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages[:20]):  # max 20 pages
            t = page.extract_text()
            if t:
                text_parts.append(f'--- Page {i+1} ---\n{t}')
    return '\n'.join(text_parts)


def _read_docx(path: Path) -> str:
    if 'docx' not in _BACKENDS:
        return '[DOCX reading requires python-docx: pip install python-docx]'
    import docx as docx_lib
    doc = docx_lib.Document(str(path))
    return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())


def _read_xlsx(path: Path) -> str:
    if 'xlsx' not in _BACKENDS:
        return '[XLSX reading requires openpyxl: pip install openpyxl]'
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    parts = []
    for sheet in wb.worksheets[:5]:  # max 5 sheets
        parts.append(f'=== Sheet: {sheet.title} ===')
        for row in sheet.iter_rows(max_row=100, values_only=True):
            if any(cell is not None for cell in row):
                parts.append('\t'.join(str(c) if c is not None else '' for c in row))
    return '\n'.join(parts)


def _read_image_ocr(path: Path) -> str:
    """Try PaddleOCR first, fallback to pytesseract, fallback to description."""
    # Try PaddleOCR (from the repo)
    try:
        from paddleocr import PaddleOCR
        ocr = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)
        result = ocr.ocr(str(path), cls=True)
        lines = []
        for line in (result[0] or []):
            if line and len(line) > 1 and line[1]:
                lines.append(line[1][0])
        return '\n'.join(lines) if lines else '[No text detected in image]'
    except Exception:
        pass

    # Try pytesseract
    try:
        import pytesseract
        from PIL import Image as PILImage
        img = PILImage.open(path)
        return pytesseract.image_to_string(img)
    except Exception:
        pass

    return f'[Image OCR unavailable. Install PaddleOCR: pip install paddleocr, or pytesseract]'


@tool(name='doc_read', description='Read any document: PDF, DOCX, XLSX, images (with OCR), or plain text files.')
async def doc_read(path: str, max_chars: int = 6000) -> dict:
    """
    Extract text from any document.
    Args:
        path: Full path to the document
        max_chars: Maximum characters to return (default 6000)
    """
    try:
        p = Path(path).expanduser().resolve()
        if not p.exists():
            return {'error': f'File not found: {path}', 'path': str(p)}

        ext = p.suffix.lower()
        content = ''

        if ext == '.pdf':
            content = _read_pdf(p)
        elif ext in ('.docx', '.doc'):
            content = _read_docx(p)
        elif ext in ('.xlsx', '.xls', '.csv'):
            if ext == '.csv':
                content = p.read_text(encoding='utf-8', errors='replace')
            else:
                content = _read_xlsx(p)
        elif ext in ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp', '.gif'):
            content = _read_image_ocr(p)
        elif ext in ('.txt', '.md', '.log', '.json', '.xml', '.html', '.py', '.js', '.ts', '.java', '.c', '.cpp'):
            content = p.read_text(encoding='utf-8', errors='replace')
        else:
            # Try as text
            try:
                content = p.read_text(encoding='utf-8', errors='replace')
            except Exception:
                return {'error': f'Unsupported file type: {ext}', 'path': str(p)}

        truncated = len(content) > max_chars
        return {
            'path':      str(p),
            'type':      ext,
            'content':   content[:max_chars],
            'chars':     len(content),
            'truncated': truncated,
        }
    except Exception as e:
        log.error(f'doc_read error: {e}')
        return {'error': str(e), 'path': path}


@tool(name='doc_extract_table', description='Extract tables from PDF or Excel documents as structured data.')
async def doc_extract_table(path: str, page: int = 1) -> dict:
    """
    Extract tables from documents.
    Args:
        path: Path to PDF or Excel file
        page: Page number for PDFs (default 1)
    """
    try:
        p = Path(path).expanduser().resolve()
        ext = p.suffix.lower()

        if ext == '.pdf' and 'pdf' in _BACKENDS:
            import pdfplumber
            with pdfplumber.open(p) as pdf:
                if page > len(pdf.pages):
                    return {'error': f'Page {page} does not exist (PDF has {len(pdf.pages)} pages)'}
                tables = pdf.pages[page - 1].extract_tables()
                return {
                    'path':   str(p),
                    'page':   page,
                    'tables': [[[str(c) if c else '' for c in row] for row in table] for table in tables],
                    'count':  len(tables),
                }

        elif ext in ('.xlsx', '.xls') and 'xlsx' in _BACKENDS:
            import openpyxl
            wb = openpyxl.load_workbook(p, data_only=True)
            ws = wb.active
            rows = []
            for row in ws.iter_rows(max_row=200, values_only=True):
                rows.append([str(c) if c is not None else '' for c in row])
            return {'path': str(p), 'table': rows, 'rows': len(rows)}

        else:
            return {'error': f'Table extraction not available for {ext}. Install pdfplumber for PDFs.'}

    except Exception as e:
        return {'error': str(e), 'path': path}


@tool(
    name='doc_save',
    description=(
        'Save edited text content back to a document file. '
        'Use this after editing a document to write it back to disk. '
        'Supports .docx (Word), .txt, .csv, .md, .html and any text format. '
        'For .docx files, creates a proper Word document. '
        'Args: path (str) — original file path, content (str) — full edited text content.'
    )
)
async def doc_save(path: str, content: str) -> dict:
    """
    Save edited content back to a document.
    - .docx / .doc  → writes a proper Word document (each paragraph on new line)
    - .txt / .md / .csv / .html / .py / etc. → plain text UTF-8
    - Always saves next to the original with '_edited' suffix to avoid overwriting
    """
    try:
        from pathlib import Path
        p = Path(path).expanduser().resolve()
        p.parent.mkdir(parents=True, exist_ok=True)

        ext = p.suffix.lower()

        # Save as proper Word document
        if ext in ('.docx', '.doc'):
            try:
                import docx as _docx
                from docx import Document
                from docx.shared import Pt
                doc = Document()
                # Preserve paragraph structure
                for para in content.split('\n'):
                    if para.strip():
                        p_obj = doc.add_paragraph(para)
                        p_obj.style.font.size = Pt(11)
                    else:
                        doc.add_paragraph('')  # blank line

                # Save as _edited.docx alongside original
                out_path = p.parent / (p.stem + '_edited.docx')
                doc.save(str(out_path))
                return {
                    'success':   True,
                    'saved_to':  str(out_path),
                    'format':    'Word Document (.docx)',
                    'note':      f'Saved as {out_path.name} (original untouched)',
                    'answer':    f'Done! Edited Word document saved to: {out_path}',
                }
            except ImportError:
                # python-docx not available — save as txt instead
                out_path = p.parent / (p.stem + '_edited.txt')
                out_path.write_text(content, encoding='utf-8')
                return {
                    'success':  True,
                    'saved_to': str(out_path),
                    'format':   'Plain Text (python-docx not installed)',
                    'answer':   f'Saved as plain text to: {out_path} (install python-docx for Word format)',
                }

        # Plain text formats
        else:
            out_path = p.parent / (p.stem + '_edited' + ext) if p.exists() else p
            out_path.write_text(content, encoding='utf-8')
            return {
                'success':  True,
                'saved_to': str(out_path),
                'format':   ext or 'text',
                'answer':   f'Done! Edited file saved to: {out_path}',
            }

    except Exception as e:
        log.error(f'doc_save error: {e}')
        return {'error': str(e), 'path': path, 'success': False}
