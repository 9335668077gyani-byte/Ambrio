# ambrio/router/tools/convert_tool.py
"""
File Conversion Tool for Ambrio.
Converts documents between formats locally, no cloud needed.

Supported conversions:
  docx  → pdf, txt
  pdf   → txt
  txt   → docx
  xlsx  → csv
  csv   → xlsx
  md    → html, txt
  html  → txt
  jpg/png → pdf
"""
import logging
from pathlib import Path
from ambrio.router.tool_registry import tool

log = logging.getLogger(__name__)

# ── Conversion matrix ─────────────────────────────────────────────────────────
_SUPPORTED = {
    '.docx': ['pdf', 'txt'],
    '.doc':  ['txt'],
    '.pdf':  ['txt'],
    '.txt':  ['docx', 'pdf'],
    '.md':   ['html', 'txt', 'docx'],
    '.html': ['txt', 'docx'],
    '.xlsx': ['csv', 'txt'],
    '.xls':  ['csv'],
    '.csv':  ['xlsx'],
    '.jpg':  ['pdf'],
    '.jpeg': ['pdf'],
    '.png':  ['pdf'],
    '.bmp':  ['pdf'],
    '.webp': ['pdf'],
    '.gif':  ['pdf'],
    '.tiff': ['pdf'],
}


def _out_path(src: Path, target_ext: str) -> Path:
    """Build output path alongside source with new extension."""
    return src.parent / (src.stem + '.' + target_ext.lstrip('.'))


@tool(
    name='doc_convert',
    description=(
        'Convert a file from one format to another. '
        'Supports: docx→pdf/txt, pdf→txt, txt→docx, xlsx→csv, csv→xlsx, '
        'md→html/txt, html→txt, jpg/png→pdf. '
        'Args: path (str) — source file path, to (str) — target format e.g. "pdf", "txt", "docx", "csv", "xlsx".'
    )
)
async def doc_convert(path: str, to: str) -> dict:
    """
    Convert a file to another format.
    Args:
        path: Absolute path to the source file
        to:   Target format — pdf, txt, docx, csv, xlsx, html
    """
    try:
        src = Path(path).expanduser().resolve()
        if not src.exists():
            return {'error': f'File not found: {path}', 'success': False}

        src_ext = src.suffix.lower()
        tgt_ext = to.lower().lstrip('.')
        out     = _out_path(src, tgt_ext)

        # ── DOCX → PDF ───────────────────────────────────────────────────────
        if src_ext in ('.docx', '.doc') and tgt_ext == 'pdf':
            try:
                from docx2pdf import convert
                convert(str(src), str(out))
                return _ok(out, 'Word → PDF via Microsoft Word')
            except Exception as e:
                # Fallback: try LibreOffice if installed
                import subprocess
                result = subprocess.run(
                    ['soffice', '--headless', '--convert-to', 'pdf',
                     '--outdir', str(out.parent), str(src)],
                    capture_output=True, timeout=30
                )
                if result.returncode == 0:
                    return _ok(out, 'Word → PDF via LibreOffice')
                return {'error': f'PDF conversion needs Microsoft Word or LibreOffice installed. '
                                 f'Install either app then retry. ({e})',
                        'success': False}

        # ── DOCX → TXT ───────────────────────────────────────────────────────
        elif src_ext in ('.docx', '.doc') and tgt_ext == 'txt':
            from docx import Document
            doc  = Document(str(src))
            text = '\n'.join(p.text for p in doc.paragraphs)
            out.write_text(text, encoding='utf-8')
            return _ok(out, 'Word → Plain Text')

        # ── PDF → TXT ────────────────────────────────────────────────────────
        elif src_ext == '.pdf' and tgt_ext == 'txt':
            try:
                import pdfplumber
                text_parts = []
                with pdfplumber.open(str(src)) as pdf:
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            text_parts.append(t)
                out.write_text('\n\n'.join(text_parts), encoding='utf-8')
                return _ok(out, 'PDF → Plain Text via pdfplumber')
            except ImportError:
                return {'error': 'Install pdfplumber: pip install pdfplumber', 'success': False}

        # ── TXT / MD → DOCX ──────────────────────────────────────────────────
        elif src_ext in ('.txt', '.md', '.html') and tgt_ext == 'docx':
            from docx import Document
            from docx.shared import Pt
            content = src.read_text(encoding='utf-8', errors='replace')
            doc = Document()
            doc.core_properties.title = src.stem
            for line in content.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line)
                    p.style.font.size = Pt(11)
                else:
                    doc.add_paragraph('')
            out = _out_path(src, 'docx')
            doc.save(str(out))
            return _ok(out, 'Text → Word Document')

        # ── TXT → PDF ────────────────────────────────────────────────────────
        elif src_ext == '.txt' and tgt_ext == 'pdf':
            # First convert to docx, then to pdf
            result = await doc_convert(path, 'docx')
            if result.get('success'):
                return await doc_convert(result['saved_to'], 'pdf')
            return result

        # ── MD → HTML ────────────────────────────────────────────────────────
        elif src_ext == '.md' and tgt_ext == 'html':
            try:
                import markdown as md_lib
                text = src.read_text(encoding='utf-8')
                html = md_lib.markdown(text, extensions=['tables', 'fenced_code'])
                full = f'<!DOCTYPE html><html><body>\n{html}\n</body></html>'
                out.write_text(full, encoding='utf-8')
                return _ok(out, 'Markdown → HTML')
            except ImportError:
                return {'error': 'Install markdown: pip install markdown', 'success': False}

        # ── MD → TXT ─────────────────────────────────────────────────────────
        elif src_ext == '.md' and tgt_ext == 'txt':
            text = src.read_text(encoding='utf-8')
            # Strip markdown symbols
            import re
            text = re.sub(r'#{1,6}\s*', '', text)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            out.write_text(text, encoding='utf-8')
            return _ok(out, 'Markdown → Plain Text')

        # ── HTML → TXT ───────────────────────────────────────────────────────
        elif src_ext == '.html' and tgt_ext == 'txt':
            try:
                from bs4 import BeautifulSoup
                html = src.read_text(encoding='utf-8')
                soup = BeautifulSoup(html, 'html.parser')
                out.write_text(soup.get_text(), encoding='utf-8')
                return _ok(out, 'HTML → Plain Text')
            except ImportError:
                return {'error': 'Install bs4: pip install beautifulsoup4', 'success': False}

        # ── XLSX → CSV ───────────────────────────────────────────────────────
        elif src_ext in ('.xlsx', '.xls') and tgt_ext == 'csv':
            import openpyxl, csv
            wb = openpyxl.load_workbook(str(src), read_only=True, data_only=True)
            ws = wb.active
            with open(str(out), 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                for row in ws.iter_rows(values_only=True):
                    writer.writerow(['' if v is None else str(v) for v in row])
            wb.close()
            return _ok(out, 'Excel → CSV')

        # ── CSV → XLSX ───────────────────────────────────────────────────────
        elif src_ext == '.csv' and tgt_ext == 'xlsx':
            import openpyxl, csv
            wb = openpyxl.Workbook()
            ws = wb.active
            with open(str(src), newline='', encoding='utf-8', errors='replace') as f:
                for row in csv.reader(f):
                    ws.append(row)
            wb.save(str(out))
            return _ok(out, 'CSV → Excel')

        # ── IMAGE → PDF ── scan-style (page = image, no white borders) ─────────
        elif src_ext in ('.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp', '.gif') and tgt_ext == 'pdf':
            try:
                from PIL import Image as PILImage

                img = PILImage.open(str(src))

                # Normalise colour mode
                if img.mode in ('RGBA', 'LA', 'PA'):
                    bg = PILImage.new('RGB', img.size, (255, 255, 255))
                    bg.paste(img.convert('RGBA'), mask=img.convert('RGBA').split()[-1])
                    img = bg
                elif img.mode == 'P':
                    img = img.convert('RGB')
                elif img.mode != 'RGB':
                    img = img.convert('RGB')

                DPI = 200

                # Upscale tiny images (WhatsApp thumbnails etc.) to min 1200px
                img_w, img_h = img.size
                long_side = max(img_w, img_h)
                if long_side < 1200:
                    scale = 1200 / long_side
                    img = img.resize(
                        (int(img_w * scale), int(img_h * scale)),
                        PILImage.LANCZOS
                    )

                # Page size = image size (scanner-style, zero white border)
                img.save(str(out), format='PDF', resolution=DPI)

                w, h = img.size
                orient = 'landscape' if w > h else 'portrait'
                return _ok(out, f'Image → PDF ({orient}, {w}×{h}px, no borders)')

            except ImportError:
                return {'error': 'Install Pillow: pip install Pillow', 'success': False}

        else:
            supported = _SUPPORTED.get(src_ext, [])
            if supported:
                return {
                    'error': f'Cannot convert {src_ext} to .{tgt_ext}. '
                             f'Supported targets for {src_ext}: {supported}',
                    'success': False
                }
            else:
                return {
                    'error': f'File type {src_ext} is not supported for conversion. '
                             f'Supported: {list(_SUPPORTED.keys())}',
                    'success': False
                }

    except Exception as e:
        log.error(f'doc_convert error: {e}')
        return {'error': str(e), 'path': path, 'success': False}


def _ok(out: Path, method: str) -> dict:
    return {
        'success':   True,
        'saved_to':  str(out),
        'method':    method,
        'file_name': out.name,
        'answer':    f'Done! Converted file saved to: {out}',
    }


@tool(
    name='doc_combine',
    description=(
        'Combine 2 images (e.g. front and back of an ID card) onto a single A4 white page PDF. '
        'Image 1 goes on the top half, image 2 on the bottom half. '
        'Perfect for Aadhaar, PAN card, driving licence — both sides on one A4 sheet. '
        'Args: path1 (str) — first image path, path2 (str) — second image path, '
        'out_name (str, optional) — output PDF filename (default: combined.pdf next to path1).'
    )
)
async def doc_combine(path1: str, path2: str, out_name: str = '') -> dict:
    """
    Combine 2 images onto one A4 white-background PDF page.
    Args:
        path1:    Absolute path to image 1 (e.g. front of ID card)
        path2:    Absolute path to image 2 (e.g. back of ID card)
        out_name: Optional output filename (e.g. "ADHAR.pdf")
    """
    try:
        from PIL import Image as PILImage

        # ── helpers ──────────────────────────────────────────────────────────
        def _load(p: str) -> 'PILImage.Image':
            img = PILImage.open(p)
            if img.mode in ('RGBA', 'LA', 'PA'):
                bg = PILImage.new('RGB', img.size, (255, 255, 255))
                bg.paste(img.convert('RGBA'), mask=img.convert('RGBA').split()[-1])
                return bg
            return img.convert('RGB')

        img1 = _load(path1)
        img2 = _load(path2)

        # ── A4 canvas at 150 DPI ─────────────────────────────────────────────
        DPI      = 150
        A4_W     = int(8.27  * DPI)   # 1240 px
        A4_H     = int(11.69 * DPI)   # 1753 px
        MARGIN   = int(0.25  * DPI)   # 37 px  (~6 mm border)
        GAP      = int(0.15  * DPI)   # 22 px  gap between cards

        canvas = PILImage.new('RGB', (A4_W, A4_H), (255, 255, 255))

        # Each image gets half the page height minus margins/gap
        slot_w = A4_W - 2 * MARGIN
        slot_h = (A4_H - 2 * MARGIN - GAP) // 2

        def _fit(img: 'PILImage.Image', max_w: int, max_h: int) -> 'PILImage.Image':
            """Scale image to fit inside max_w × max_h keeping aspect ratio."""
            iw, ih = img.size
            scale  = min(max_w / iw, max_h / ih)
            new_w  = int(iw * scale)
            new_h  = int(ih * scale)
            return img.resize((new_w, new_h), PILImage.LANCZOS)

        # ── Place image 1 — top slot ─────────────────────────────────────────
        img1_fit = _fit(img1, slot_w, slot_h)
        x1 = MARGIN + (slot_w - img1_fit.width)  // 2
        y1 = MARGIN + (slot_h - img1_fit.height) // 2
        canvas.paste(img1_fit, (x1, y1))

        # ── Place image 2 — bottom slot ──────────────────────────────────────
        img2_fit = _fit(img2, slot_w, slot_h)
        x2 = MARGIN + (slot_w - img2_fit.width)  // 2
        y2 = MARGIN + slot_h + GAP + (slot_h - img2_fit.height) // 2
        canvas.paste(img2_fit, (x2, y2))

        # ── Save ─────────────────────────────────────────────────────────────
        p1 = Path(path1)
        if out_name:
            out = p1.parent / out_name
            if not out.suffix:
                out = out.with_suffix('.pdf')
        else:
            out = p1.parent / (p1.stem + '_combined.pdf')

        canvas.save(str(out), format='PDF', resolution=DPI)

        return {
            'success':   True,
            'saved_to':  str(out),
            'file_name': out.name,
            'answer':    f'✅ Combined both images on A4 → saved to: {out}',
        }

    except ImportError:
        return {'error': 'Install Pillow: pip install Pillow', 'success': False}
    except Exception as e:
        log.error(f'doc_combine error: {e}')
        return {'error': str(e), 'success': False}
